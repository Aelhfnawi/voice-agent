"""
Document extraction module for the Technical RAG Assistant.

This module handles extraction of text from various document formats:
- PDF files (using PyPDF2 and pdfplumber as fallback)
- DOCX files (using python-docx)
- TXT files (direct reading)

The extractor intelligently handles encoding issues and extraction failures.
"""

import logging
from pathlib import Path
from typing import Union, Optional
import io

# PDF extraction libraries
try:
    import PyPDF2
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logging.warning("PyPDF2 not installed. PDF extraction will be limited.")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not installed. Fallback PDF extraction unavailable.")

# DOCX extraction
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. DOCX extraction unavailable.")


logger = logging.getLogger(__name__)


class DocumentExtractor:
    """
    Extracts text content from PDF, DOCX, and TXT files.
    
    Supports multiple extraction methods with automatic fallback.
    """
    
    def __init__(self, extraction_method: str = 'auto'):
        """
        Initialize the DocumentExtractor.
        
        Args:
            extraction_method: Preferred extraction method for PDFs.
                Options: 'auto', 'pypdf', 'pdfplumber'
                'auto' tries PyPDF2 first, falls back to pdfplumber if needed
        """
        self.extraction_method = extraction_method
        logger.info(f"DocumentExtractor initialized with method: {extraction_method}")
    
    def extract(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content as a string
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
            Exception: If extraction fails
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        extension = file_path.suffix.lower()
        
        logger.info(f"Extracting text from: {file_path.name}")
        
        try:
            # Route to appropriate extraction method
            if extension == '.pdf':
                text = self._extract_pdf(file_path)
            elif extension == '.docx':
                text = self._extract_docx(file_path)
            elif extension == '.txt':
                text = self._extract_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
            
            if not text or not text.strip():
                logger.warning(f"No text extracted from {file_path.name}")
                return ""
            
            logger.info(f"Successfully extracted {len(text)} characters from {file_path.name}")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting from {file_path.name}: {e}")
            raise
    
    def _extract_pdf(self, file_path: Path) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        # Try primary method
        if self.extraction_method in ['auto', 'pypdf'] and PYPDF_AVAILABLE:
            try:
                text = self._extract_pdf_pypdf(file_path)
                if text and text.strip():
                    return text
                logger.warning(f"PyPDF2 extracted empty text from {file_path.name}")
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed for {file_path.name}: {e}")
        
        # Try fallback method
        if self.extraction_method in ['auto', 'pdfplumber'] and PDFPLUMBER_AVAILABLE:
            try:
                text = self._extract_pdf_pdfplumber(file_path)
                if text and text.strip():
                    return text
                logger.warning(f"pdfplumber extracted empty text from {file_path.name}")
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed for {file_path.name}: {e}")
        
        raise Exception(f"All PDF extraction methods failed for {file_path.name}")
    
    def _extract_pdf_pypdf(self, file_path: Path) -> str:
        """
        Extract text from PDF using PyPDF2.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            logger.debug(f"Processing {num_pages} pages with PyPDF2")
            
            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text:
                        text_content.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    continue
        
        return "\n\n".join(text_content)
    
    def _extract_pdf_pdfplumber(self, file_path: Path) -> str:
        """
        Extract text from PDF using pdfplumber (better for complex layouts).
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages)
            
            logger.debug(f"Processing {num_pages} pages with pdfplumber")
            
            for page_num, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    
                    if text:
                        text_content.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    continue
        
        return "\n\n".join(text_content)
    
    def _extract_docx(self, file_path: Path) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX extraction. Install it with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            logger.debug(f"Extracted {len(text_content)} text blocks from DOCX")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise
    
    def _extract_txt(self, file_path: Path) -> str:
        """
        Extract text from a plain text file.
        
        Handles multiple encodings automatically.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Extracted text content
        """
        # Try common encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    logger.debug(f"Successfully read TXT file with {encoding} encoding")
                    return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error reading TXT file with {encoding}: {e}")
                continue
        
        raise Exception(f"Failed to read {file_path.name} with any supported encoding")
    
    def extract_batch(self, file_paths: list[Union[str, Path]]) -> dict[str, str]:
        """
        Extract text from multiple files.
        
        Args:
            file_paths: List of file paths to extract from
            
        Returns:
            Dictionary mapping file paths to extracted text
            Failed extractions will have empty string values
        """
        results = {}
        
        for file_path in file_paths:
            try:
                text = self.extract(file_path)
                results[str(file_path)] = text
            except Exception as e:
                logger.error(f"Batch extraction failed for {file_path}: {e}")
                results[str(file_path)] = ""
        
        return results
    
    def get_document_metadata(self, file_path: Union[str, Path]) -> dict:
        """
        Extract metadata from a document.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing metadata
        """
        file_path = Path(file_path)
        
        metadata = {
            'filename': file_path.name,
            'extension': file_path.suffix.lower(),
            'size_bytes': file_path.stat().st_size,
            'size_mb': round(file_path.stat().st_size / (1024 * 1024), 2)
        }
        
        # Add PDF-specific metadata
        if file_path.suffix.lower() == '.pdf' and PYPDF_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['num_pages'] = len(pdf_reader.pages)
                    
                    # Extract PDF metadata if available
                    if pdf_reader.metadata:
                        metadata['title'] = pdf_reader.metadata.get('/Title', '')
                        metadata['author'] = pdf_reader.metadata.get('/Author', '')
                        metadata['subject'] = pdf_reader.metadata.get('/Subject', '')
            except Exception as e:
                logger.warning(f"Could not extract PDF metadata: {e}")
        
        return metadata


if __name__ == "__main__":
    # Test the extractor
    import sys
    
    logging.basicConfig(level=logging.INFO)
    
    if len(sys.argv) < 2:
        print("Usage: python extractor.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    extractor = DocumentExtractor()
    
    try:
        # Extract metadata
        metadata = extractor.get_document_metadata(file_path)
        print(f"\n{'='*80}")
        print(f"METADATA: {metadata['filename']}")
        print(f"{'='*80}")
        for key, value in metadata.items():
            print(f"{key}: {value}")
        
        # Extract text
        text = extractor.extract(file_path)
        print(f"\n{'='*80}")
        print(f"EXTRACTED TEXT ({len(text)} characters)")
        print(f"{'='*80}")
        print(text[:500] + "..." if len(text) > 500 else text)
        print(f"\n{'='*80}\n")
        
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)